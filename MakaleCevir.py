# Import libraries
from PIL import Image
import pytesseract
import sys
from pdf2image import convert_from_path
import os
import translate
import re
from time import sleep
import docx
from docx import Document


def write_to_file(text,translated_text):
    #If document already exists
    if os.path.exists(outfile+str(i)+'.docx'):
        doc = Document(outfile+str(i)+'.docx')
        # for para in doc.paragraphs:
        #     doc.add_paragraph(para)
        try:
            doc.add_paragraph(text)
            doc.add_paragraph("+++++++++++++++++++++++++++++++++++")
            doc.add_paragraph(translated_text)
            doc.add_paragraph("\n\n")
        except Exception as e:
            print("Dosya yazma işlemi hatası\n")
            print(e)
        doc.save(outfile+str(i)+'.docx')

    else:
        # create document object
        doc = docx.Document()
        try:
            doc.add_paragraph(text)
            doc.add_paragraph("+++++++++++++++++++++++++++++++++++")
            doc.add_paragraph(translated_text)
            doc.add_paragraph("\n\n")
        except Exception as e:
            print("Dosya yazma işlemi hatası\n")
            print(e)
        doc.save(outfile+str(i)+'.docx')

    # write to .txt file.
    # f = open(outfile+str(i)+".txt", "a")
    # try:
    #     f.write(text)
    #     f.write("\n++++++++++++++++++++++++++\n")
    #     f.write(translated_text)
    #     f.write("\n\n\n\n")
    # except Exception as e:
    #     print("Dosya yazma işlemi hatası\n")
    #     print(e)
    # f.close()
    # print(f"Sayfa {i} çeviri işlemi devam ediyor.")

    # Close the file after writing all the text.

def parantez_icinde_alinti_temizleme(text):
    element_to_delete = []
    x = re.findall("\([^)]*\)",text) # parantez içinde text var
    #print(x)
    if x != None:
        for element in x:
            if ',' in element: # parantez içindeki text , içeriyor
                element_to_delete.append(element)
                #print(text)
            else:
                pass
                #print("False")
        #print(element_to_delete)
    for eleman in element_to_delete:
        text = text.replace(eleman, '')

    return text #succeed.


def metin_birlestirme(text):
    if i == 1 and "abstract" in text.lower():
        #print("ABSTRACT in TEXT")
        indis = text.lower().index('abstract')
        #print(indis)
        text = text[indis:]
        #print(text)

    text_list = text.split('\n\n')
    returning_text = ""
    el2 = ""
    for el in text_list:
        if '\n' in el:
            el = el.replace('\n',' ')
        #print("el = ",el)
        if len(el) < 5000:
            #print("Buradasin1")
            try:
                translated_text = translate.translate(el)
            except:
                pass
            else:
                write_to_file(el,translated_text)
        else:
            #print("Buradasin2")
            el_list = el.split('.')
            sep_text = ""
            part_translated_text = ""
            j = 0
            while True:
                #print("Buradasin3")
                while len(sep_text)<4000:
                    try:
                        sep_text = sep_text + el_list[j]
                        j += 1
                    except:
                        break
                part_translated_text = part_translated_text + translate.translate(sep_text)
                if sep_text == "":
                    write_to_file(el,part_translated_text)
                    break
                sep_text = ""


PDF_file = str(input("Çevrilecek makale PDF dosyasının adını uzantısı ile birlikte yazınız.(örnek: Data_Mining.pdf)\n"))

'''
Part #1 : Converting PDF to images
'''
print("PDF Sayfaları İşlenebilir Resim Formatına Dönüştürülüyor. Bu İşlemin Süresi Bilgisayarınızın Hızına ve PDF Uzunluğuna Göre Değişir.")
print("Genelde 10 Sayfalık bir PDF 1 dakikadan kısa sürer.")
try:
    path = os.getcwd()
    PDF_file = path + '\\' + PDF_file
    # Store all the pages of the PDF in a variable
    poppler_path = path + r"\assets\poppler-22.01.0\Library\bin"
    #print("poppler_path ", poppler_path)
    pages = convert_from_path(PDF_file, 500, poppler_path = poppler_path)

except Exception as er:
    print("Dosya bulunamadı. Lütfen program ile dosyayı aynı dizine koyup dosyanın adını uzantısı ile beraber girin.")
    print(er)
    sleep(100000)
# Counter to store images of each page of PDF to image
image_counter = 1

# Iterate through all the pages stored above
for page in pages:
    # Declaring filename for each page of PDF as JPG
    # For each page, filename will be:
    # PDF page 1 -> page_1.jpg
    # PDF page 2 -> page_2.jpg
    # PDF page 3 -> page_3.jpg
    # ....
    # PDF page n -> page_n.jpg
    filename = "page_"+str(image_counter)+".jpg"

    # Save the image of the page in system
    page.save(filename, 'JPEG')

    # Increment the counter to update filename
    image_counter = image_counter + 1

'''
Part #2 - Recognizing text from the images using OCR
'''

# Variable to get count of total number of pages
filelimit = image_counter-1

# Creating a text file to write the output
outfile = "Sayfa "

# Open the file in append mode so that
# All contents of all images are added to the same file


# Iterate from 1 to total number of pages
for i in range(1, filelimit + 1):
    print(f"Sayfa {i} için çeviri işlemi yapılıyor.")
    # Set filename to recognize text from
    # Again, these files will be:
    # page_1.jpg
    # page_2.jpg
    # ....
    # page_n.jpg
    filename = path + "\\page_"+str(i)+".jpg"

    #text = f"START OF PAGE {i} \n\n"
    #print(text)
    #print("---------------11111111111111111111111111-------------------------")
    # Recognize the text as string in image using pytesserct
    try:
        pytesseract.pytesseract.tesseract_cmd = path + r'\assets\Tesseract-OCR\tesseract.exe'
        text = str(((pytesseract.image_to_string(Image.open(filename)))))
    except Exception as e:
        print("Resimden Metine Geçiş Hatası")
        print(e)
        sleep(100000)
    #print(text)
    #print("---------------22222222222222222222222222-------------------------")
    # The recognized text is stored in variable text
    # Any string processing may be applied on text
    # Here, basic formatting has been done:
    # In many PDFs, at line ending, if a word can't
    # be written fully, a 'hyphen' is added.
    # The rest of the word is written in the next line
    # Eg: This is a sample text this word here GeeksF-
    # orGeeks is half on first line, remaining on next.
    # To remove this, we replace every '-\n' to ''.
    text = text.replace('-\n', '')
    #text = text + f"\n\n END OF PAGE {i}"
    #print(text)
    if "REFERENCES" in text:
        referances_index = text.index("REFERENCES")
        text = text[0:referances_index]
        text = parantez_icinde_alinti_temizleme(text)
        text = metin_birlestirme(text)
        break



    text = parantez_icinde_alinti_temizleme(text)
    text = metin_birlestirme(text)
    #write_to_file(translated_text)
print("Çeviri Tamamlandı. 10 Saniye sonra kendimi kapatacağım.")
sleep(10)
